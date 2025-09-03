#!/usr/bin/env python

import docx
import os

# `cd` to directory with *.pdf
cwd = os.getcwd() + '/resources'
os.chdir(cwd)

# open docx object
doc = docx.Document('demo.docx')

# first paragraph text
doc.paragraphs[0].text

# second paragraph runs (style)
par = doc.paragraphs[1]
par.runs[0].text                                        # 'A plain paragraph with'

# check runs style
par.runs[0].bold == None                                # True

# formatting
par.runs[3].text
par.runs[3].underline = True                            # underline
par.runs[3].text = 'italic and underlined.'

# write to file
doc.save('demo2.docx')

# more formatting
par.style = 'Title'
doc.save('demo3.docx')                                  # sentence now in title style (massive font)

# create docx
doc = docx.Document()                                   # in-memory new docx

# add paragraph
doc.add_paragraph('Hello. This is a paragraph. Ish. ')
doc.add_paragraph('Another paragraph. Kinda. ')

# write to file
doc.save('demo4.docx')

# add sentence (paragraph)
par = doc.paragraphs[0]
par.add_run('This is a new run. ')
par.runs
par.runs[1].bold = True                                 # bold 'This is a new run. '

# write to file
doc.save('demo5.docx')

# scrape all text from docx
def get_text(fn):
    doc = docx.Document(fn)
    full_txt = []
    for par in doc.paragraphs:
        full_txt.append(par.text)
    return '\n'.join(full_txt)

print(get_text('demo.docx'))
